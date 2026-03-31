import pandas as pd
from docxtpl import DocxTemplate
from docx2pdf import convert
from datetime import datetime
import os
import json
import msal
import requests
import base64
from dotenv import load_dotenv

# --- ASSET GEN LOGIC ---
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

# --- PDF OVERLAY ---
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Load environment variables from .env file
load_dotenv()

# ──────────────────────────────────────────────
# SHAREPOINT CONFIGURATION
# ──────────────────────────────────────────────
SHAREPOINT_TENANT_ID = os.getenv("SHAREPOINT_TENANT_ID")
SHAREPOINT_CLIENT_ID = os.getenv("SHAREPOINT_CLIENT_ID")
SHAREPOINT_CLIENT_SECRET = os.getenv("SHAREPOINT_CLIENT_SECRET")
SHAREPOINT_FILE_URL = os.getenv("SHAREPOINT_FILE_URL")

# ──────────────────────────────────────────────
# OVERLAY CALIBRATION
# ✅ MEASURED from actual PDF border lines (pdfplumber rect detection)
#    Columns confirmed: 7 equal-width cols across 12.7mm → 197.19mm
#    Row height: ~16.05mm (extremely consistent across all 7 rows)
# ──────────────────────────────────────────────
OVERLAY_CONFIG = {
    "page_size": "A4",
    # ── PAGE 1 ──────────────────────────────────────────────────────────────
    "table_data_start_y_mm": 156.700,  # top of data row 1 (border line exact)
    "rows_per_page1": 7,  # 7 data rows fit on page 1
    # ── PAGE 2 ──────────────────────────────────────────────────────────────
    "table_data_start_y_mm_page2": 61.600,
    # ── ROW HEIGHT ──────────────────────────────────────────────────────────
    "row_height_mm": 16.055,  # measured from border rects, ultra consistent
    # ── COLUMN POSITIONS (from actual PDF border lines) ──────────────────────
    # x0 = left cell border (mm from page left edge)
    # 7 cols: Asset ID | Item | Brand | Model | Serial No. | Notes | Sign
    "col_x_mm": [12.700, 39.000, 65.250, 91.540, 117.840, 144.590, 170.890],
    "col_w_mm": [26.300, 26.250, 26.290, 26.300, 26.750, 26.300, 26.300],
    "font_size": 9,
}


# ──────────────────────────────────────────────
# FONT SETUP
# ReportLab cannot use system fonts by name — TTF files must be registered.
# Set OVERLAY_FONT_NAME to any font you want. If the TTF is not found,
# the script automatically falls back to built-in Helvetica.
# ──────────────────────────────────────────────
OVERLAY_FONT_NAME = "Cambria"  # ← change this to any font name you want

# Candidate TTF paths for Windows, macOS, Linux
# Font candidates: list of (path, subfont_index)
# .ttc = TrueType Collection (multiple fonts in one file) — specify index
# .ttf = single font — index always 0
_FONT_CANDIDATES = {
    "Cambria": [
        ("C:/Windows/Fonts/cambria.ttc", 0),  # .ttc index 0 = Cambria Regular
        ("C:/Windows/Fonts/Cambria.ttc", 0),
        (
            os.path.expanduser("~")
            + "/AppData/Local/Microsoft/Windows/Fonts/cambria.ttc",
            0,
        ),
        ("C:/Windows/Fonts/cambria.ttf", 0),  # fallback in case .ttf exists
    ],
    "CambriaBold": [
        ("C:/Windows/Fonts/cambriab.ttf", 0),
        ("C:/Windows/Fonts/CambriaBold.ttf", 0),
    ],
    "CambriaItalic": [
        ("C:/Windows/Fonts/cambriai.ttf", 0),
    ],
    "Arial": [
        ("C:/Windows/Fonts/arial.ttf", 0),
        ("/Library/Fonts/Arial.ttf", 0),
        ("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf", 0),
    ],
    "Calibri": [
        ("C:/Windows/Fonts/calibri.ttf", 0),
        (
            os.path.expanduser("~")
            + "/AppData/Local/Microsoft/Windows/Fonts/calibri.ttf",
            0,
        ),
    ],
    "Times-Roman": [],  # built-in
    "Helvetica": [],  # built-in
}


def _register_font(font_name):
    """
    Registers a TTF or TTC font with ReportLab.
    .ttc files are handled via subfontIndex=0 (regular weight).
    Falls back to Helvetica if the font file cannot be found.
    """
    # Built-in PostScript fonts — no registration needed
    if font_name in (
        "Helvetica",
        "Helvetica-Bold",
        "Helvetica-Oblique",
        "Times-Roman",
        "Times-Bold",
        "Courier",
        "Courier-Bold",
    ):
        return font_name

    # Already registered in this session?
    try:
        pdfmetrics.getFont(font_name)
        return font_name
    except Exception:
        pass

    import glob as _glob

    candidates = list(_FONT_CANDIDATES.get(font_name, []))

    # Auto-discover .ttf and .ttc files in common font directories
    for font_dir in [
        "C:/Windows/Fonts",
        os.path.expanduser("~") + "/AppData/Local/Microsoft/Windows/Fonts",
        "/Library/Fonts",
        "/System/Library/Fonts",
        "/usr/share/fonts/truetype",
    ]:
        for ext in ["ttf", "ttc", "TTF", "TTC"]:
            for path in _glob.glob(f"{font_dir}/{font_name}*.{ext}"):
                candidates.append((path, 0))
            for path in _glob.glob(f"{font_dir}/{font_name.lower()}*.{ext}"):
                candidates.append((path, 0))

    for path, idx in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path, subfontIndex=idx))
                ext = os.path.splitext(path)[1].upper()
                print(
                    f"  [Font] Registered '{font_name}' from: {os.path.basename(path)}"
                    f"{f' (index {idx})' if ext == '.TTC' else ''}"
                )
                return font_name
            except Exception as e:
                print(f"  [Font] Could not load {os.path.basename(path)}: {e}")

    print(f"  [Font] '{font_name}' not found — falling back to Helvetica.")
    return "Helvetica"


# Register at import time; resolved name is what all overlay functions use
OVERLAY_FONT = _register_font(OVERLAY_FONT_NAME)


def read_calibration_from_pdf(pdf_path):
    """
    Reads exact column and row positions from a generated form PDF
    by detecting the thin border rectangles drawn by the table.
    Returns a dict ready to use as OVERLAY_CONFIG geometry fields,
    or None on failure.
    """
    try:
        import pdfplumber
    except ImportError:
        print("  [Error] pdfplumber required: pip install pdfplumber")
        return None

    results = {}
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            pt = lambda mm: mm * 72 / 25.4
            mm = lambda p: p * 25.4 / 72

            # Search full page height for table rects
            rects = page.rects
            border_x = set()
            border_y = set()
            for r in rects:
                w = mm(r["x1"] - r["x0"])
                h = mm(r["bottom"] - r["top"])
                if w < 0.5 and h > 3.0:
                    border_x.add(round(mm(r["x0"]), 2))
                if h < 0.5 and w > 3.0:
                    border_y.add(round(mm(r["top"]), 2))

            border_x = sorted(border_x)
            border_y = sorted(border_y)

            if len(border_x) < 2 or len(border_y) < 2:
                continue

            col_x0 = border_x[:-1]
            col_w = [border_x[i + 1] - border_x[i] for i in range(len(border_x) - 1)]
            row_ys = border_y  # includes header top

            # Header = row 0, first data row starts at row_ys[1]
            data_start_y = row_ys[1] if len(row_ys) > 1 else row_ys[0]
            data_row_ys = row_ys[1:]
            row_heights = [
                data_row_ys[i + 1] - data_row_ys[i] for i in range(len(data_row_ys) - 1)
            ]
            avg_row_h = sum(row_heights) / len(row_heights) if row_heights else 16.0
            num_data_rows = len(row_heights)

            results[page_num] = {
                "data_start_y": data_start_y,
                "avg_row_h": avg_row_h,
                "num_data_rows": num_data_rows,
                "col_x0": col_x0,
                "col_w": col_w,
            }

    return results if results else None


PRINT_LOG_FILE = "print_log.json"

# ══════════════════════════════════════════════
# UTILITIES
# ══════════════════════════════════════════════


def separator(char="─", width=60):
    print(char * width)


def header(title):
    separator("═")
    print(f"  {title}")
    separator("═")


def section(title):
    print()
    separator()
    print(f"  {title}")
    separator()


def ask(prompt, valid=None, default=None):
    """Prompt the user, optionally validate against a list, support default."""
    while True:
        suffix = ""
        if valid:
            suffix = f" [{'/'.join(valid)}]"
        if default is not None:
            suffix += f" (default: {default})"
        raw = input(f"  {prompt}{suffix}: ").strip()
        if raw == "" and default is not None:
            return default
        if valid and raw.lower() not in [v.lower() for v in valid]:
            print(f"  [!] Please enter one of: {', '.join(valid)}")
            continue
        return raw


def ask_int(prompt, min_val=0, max_val=9999, default=None):
    """Prompt for an integer within range."""
    while True:
        suffix = f" ({min_val}–{max_val})"
        if default is not None:
            suffix += f" (default: {default})"
        raw = input(f"  {prompt}{suffix}: ").strip()
        if raw == "" and default is not None:
            return default
        try:
            val = int(raw)
            if min_val <= val <= max_val:
                return val
            print(f"  [!] Enter a number between {min_val} and {max_val}.")
        except ValueError:
            print("  [!] Invalid number.")


# ══════════════════════════════════════════════
# PRINT LOG
# ══════════════════════════════════════════════


def load_print_log():
    if os.path.exists(PRINT_LOG_FILE):
        with open(PRINT_LOG_FILE, "r") as f:
            return json.load(f)
    return {}


def save_print_log(log):
    with open(PRINT_LOG_FILE, "w") as f:
        json.dump(log, f, indent=2)


def log_key(emp_id, doc_type):
    return f"{emp_id}_{doc_type}"


def get_printed_asset_ids(emp_id, doc_type):
    log = load_print_log()
    return set(log.get(log_key(emp_id, doc_type), {}).get("printed_ids", []))


def mark_assets_printed(emp_id, doc_type, asset_ids):
    log = load_print_log()
    k = log_key(emp_id, doc_type)
    if k not in log:
        log[k] = {"printed_ids": [], "history": []}
    existing = set(log[k]["printed_ids"])
    new_ids = [str(i) for i in asset_ids if str(i) not in existing]
    log[k]["printed_ids"].extend(new_ids)
    log[k]["history"].append(
        {
            "timestamp": datetime.now().isoformat(),
            "added_ids": new_ids,
        }
    )
    save_print_log(log)
    print(f"  [+] Logged {len(new_ids)} asset(s) as printed.")


def clear_print_log(emp_id, doc_type):
    log = load_print_log()
    k = log_key(emp_id, doc_type)
    if k in log:
        del log[k]
        save_print_log(log)
        print(f"  [+] Print log cleared for {emp_id} / {doc_type}.")
    else:
        print(f"  [!] No log found for {emp_id} / {doc_type}.")


def show_print_log(emp_id, doc_type):
    log = load_print_log()
    k = log_key(emp_id, doc_type)
    if k not in log or not log[k]["printed_ids"]:
        print("  [!] No print history found.")
        return
    print(f"\n  Previously printed asset IDs:")
    for pid in log[k]["printed_ids"]:
        print(f"    • {pid}")
    if log[k].get("history"):
        print(f"\n  Print sessions:")
        for h in log[k]["history"]:
            ids = ", ".join(h["added_ids"]) if h["added_ids"] else "(none new)"
            print(f"    [{h['timestamp'][:19]}]  {ids}")


# ══════════════════════════════════════════════
# SHAREPOINT
# ══════════════════════════════════════════════


def get_sharepoint_data(file_url, tenant_id, client_id, client_secret, sheet_names):
    print("  [-] Connecting to SharePoint...")
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    app = msal.ConfidentialClientApplication(
        client_id, authority=authority, client_credential=client_secret
    )
    token_result = app.acquire_token_for_client(
        scopes=["https://graph.microsoft.com/.default"]
    )
    if "access_token" not in token_result:
        raise Exception(
            f"Token acquisition failed: {token_result.get('error_description', token_result.get('error'))}"
        )
    headers = {"Authorization": f"Bearer {token_result['access_token']}"}
    share_token = "u!" + base64.urlsafe_b64encode(file_url.encode()).decode().rstrip(
        "="
    )
    item_resp = requests.get(
        f"https://graph.microsoft.com/v1.0/shares/{share_token}/driveItem",
        headers=headers,
    )
    item_resp.raise_for_status()
    drive_item = item_resp.json()
    item_id = drive_item["id"]
    drive_id = drive_item["parentReference"]["driveId"]
    print(f"  [-] File resolved: {drive_item['name']}")
    base_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets"
    dataframes = {}
    for sheet in sheet_names:
        resp = requests.get(f"{base_url}/{sheet}/usedRange", headers=headers)
        resp.raise_for_status()
        values = resp.json()["values"]
        dataframes[sheet] = pd.DataFrame(values[1:], columns=values[0])
        print(f"  [-] Sheet '{sheet}': {len(values) - 1} rows loaded.")
    return dataframes


# ══════════════════════════════════════════════
# DOCX TABLE HELPERS
# ══════════════════════════════════════════════


def format_table_style(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def keep_row_together(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "true")
    trPr.append(cantSplit)


def set_row_height(row, height=800):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def set_cell_vertical_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def clean_text(val):
    if pd.isna(val):
        return ""
    return str(val).strip()


# ══════════════════════════════════════════════
# CALIBRATION TOOL
# ══════════════════════════════════════════════


def recalibrate_from_pdf(pdf_path):
    """
    Auto-measures OVERLAY_CONFIG values from a PDF using border rect detection
    and prints the exact block to paste into the script.
    """
    print(f"\n  Reading: {pdf_path}")
    calib = read_calibration_from_pdf(pdf_path)
    if not calib:
        print("  [Error] Could not detect table borders in PDF.")
        return

    p1 = calib.get(1, {})
    p2 = calib.get(2, {})

    for page_num, p in calib.items():
        print(
            f"\n  Page {page_num}:  data_start_y={p['data_start_y']:.3f}mm  "
            f"row_h={p['avg_row_h']:.3f}mm  "
            f"{p['num_data_rows']} data rows  {len(p['col_x0'])} cols"
        )
        for i, (x0, w) in enumerate(zip(p["col_x0"], p["col_w"])):
            names = [
                "Asset ID",
                "Item",
                "Brand",
                "Model",
                "Serial No.",
                "Notes",
                "Sign",
            ]
            n = names[i] if i < len(names) else f"col{i}"
            print(
                f"    col {i} ({n:<12}): x0={x0:.3f}mm  w={w:.3f}mm  center={(x0+w/2):.3f}mm"
            )

    print("\n")
    print("  ╔══════════════════════════════════════════════════════════╗")
    print("  ║  PASTE THIS INTO main.py → OVERLAY_CONFIG               ║")
    print("  ╚══════════════════════════════════════════════════════════╝")
    print()
    print("OVERLAY_CONFIG = {")
    print(f'    "page_size": "A4",')
    print(f'    "table_data_start_y_mm":       {p1.get("data_start_y",0):.3f},')
    print(f'    "rows_per_page1":               {p1.get("num_data_rows",7)},')
    if p2:
        print(f'    "table_data_start_y_mm_page2": {p2.get("data_start_y",0):.3f},')
    print(f'    "row_height_mm":                {p1.get("avg_row_h",16):.3f},')
    print(f'    "col_x_mm": {[round(x,3) for x in p1.get("col_x0",[])]},')
    print(f'    "col_w_mm": {[round(w,3) for w in p1.get("col_w",[])]},')
    print(f'    "font_size": 9,')
    print("}")
    print()

    """
    Generates a ruler grid PDF.
    Print on blank paper → hold over your printed form → read exact mm values.
    """
    cfg = OVERLAY_CONFIG
    pagesize = A4 if cfg["page_size"] == "A4" else (215.9 * mm, 279.4 * mm)
    page_w, page_h = pagesize

    c = rl_canvas.Canvas(out_path, pagesize=pagesize)
    c.setFont("Helvetica", 5)

    # Horizontal lines every 1 mm, labelled every 5 mm
    c.setStrokeColorRGB(0.85, 0.85, 0.85)
    c.setLineWidth(0.3)
    for y_mm in range(0, 300, 1):
        y = page_h - y_mm * mm
        c.line(0, y, page_w, y)
        if y_mm % 5 == 0:
            c.setStrokeColorRGB(0.6, 0.6, 0.6)
            c.setLineWidth(0.6)
            c.line(0, y, page_w, y)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(1 * mm, y + 0.5 * mm, f"{y_mm}")
            c.setStrokeColorRGB(0.85, 0.85, 0.85)
            c.setLineWidth(0.3)

    # Vertical lines every 5 mm
    for x_mm in range(0, 220, 5):
        x = x_mm * mm
        c.line(x, 0, x, page_h)
        c.drawString(x + 0.3 * mm, page_h - 5 * mm, f"{x_mm}")

    # Draw current column positions from config (in blue)
    c.setStrokeColorRGB(0.1, 0.4, 0.9)
    c.setLineWidth(1.0)
    for x_mm in cfg["col_x_mm"]:
        x = x_mm * mm
        c.line(x, 0, x, page_h)

    # Draw current table_data_start_y (in red)
    c.setStrokeColorRGB(0.9, 0.1, 0.1)
    c.setLineWidth(1.0)
    y_start = page_h - cfg["table_data_start_y_mm"] * mm
    c.line(0, y_start, page_w, y_start)

    # Draw a few row lines in green
    c.setStrokeColorRGB(0.1, 0.7, 0.2)
    c.setLineWidth(0.5)
    for i in range(1, 10):
        y = page_h - (cfg["table_data_start_y_mm"] + i * cfg["row_height_mm"]) * mm
        c.line(0, y, page_w, y)

    c.save()
    print(f"\n  [+] Calibration grid saved → {out_path}")
    print("      BLUE lines  = current column X positions")
    print("      RED line    = current table_data_start_y_mm")
    print("      GREEN lines = subsequent rows")
    print("      Adjust OVERLAY_CONFIG values to match your printed template.\n")
    os.startfile(out_path)


# ══════════════════════════════════════════════
# PDF OVERLAY ENGINE
# ══════════════════════════════════════════════

# ══════════════════════════════════════════════
# TEXT WRAPPING & NOTES CELL HELPERS
# ══════════════════════════════════════════════


def wrap_text(text, font, size, max_width_mm, canvas_obj):
    """Word-wraps text into lines that fit within max_width_mm."""
    words = text.split()
    if not words:
        return []
    lines, current = [], ""
    for word in words:
        test = (current + " " + word).strip()
        if canvas_obj.stringWidth(test, font, size) * 25.4 / 72 <= max_width_mm:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def best_fit_note(text, font, cell_w_mm, cell_h_mm, padding_mm, canvas_obj):
    """
    Returns (font_size, [lines]) - largest font where all wrapped lines
    fit vertically inside the cell. Falls back to 6pt with truncation.
    """
    usable_w = cell_w_mm - 2 * padding_mm
    usable_h = cell_h_mm - 2 * padding_mm
    for fsize in [9, 8, 7, 6]:
        line_h_mm = fsize * 1.2 * 25.4 / 72
        lines = wrap_text(text, font, fsize, usable_w, canvas_obj)
        if len(lines) * line_h_mm <= usable_h:
            return fsize, lines
    # Hard truncate at 6pt
    fsize = 6
    line_h_mm = fsize * 1.2 * 25.4 / 72
    max_lines = int(usable_h / line_h_mm)
    lines = wrap_text(text, font, fsize, usable_w, canvas_obj)
    return fsize, lines[:max_lines]


def draw_notes_in_cell(
    canvas_obj,
    note_text,
    cell_x0_mm,
    cell_y_top_mm,
    cell_w_mm,
    cell_h_mm,
    page_h_pt,
    padding_mm=1.5,
):
    """
    Draws wrapped, auto-sized note text vertically centered inside a cell.
    All coordinates in mm from top-left of page.
    """
    if not note_text.strip():
        return

    fsize, lines = best_fit_note(
        note_text, OVERLAY_FONT, cell_w_mm, cell_h_mm, padding_mm, canvas_obj
    )

    line_h_mm = fsize * 1.2 * 25.4 / 72
    block_h_mm = len(lines) * line_h_mm
    top_offset_mm = (cell_h_mm - block_h_mm) / 2  # vertical center

    canvas_obj.setFont(OVERLAY_FONT, fsize)
    usable_w = cell_w_mm - 2 * padding_mm

    for i, line in enumerate(lines):
        line_top_mm = cell_y_top_mm + top_offset_mm + i * line_h_mm
        y_pt = page_h_pt - (line_top_mm + line_h_mm * 0.75) * mm
        text_w_mm = canvas_obj.stringWidth(line, OVERLAY_FONT, fsize) * 25.4 / 72
        x_mm = cell_x0_mm + padding_mm + (usable_w - text_w_mm) / 2
        canvas_obj.drawString(x_mm * mm, y_pt, line)

    canvas_obj.setFont(OVERLAY_FONT, OVERLAY_CONFIG["font_size"])


def generate_overlay_pdf(rows_to_print, out_path, calib=None):
    """
    rows_to_print : list of dicts  {"target_row": int (1-indexed), "page": int, "values": [7 strings]}
    calib         : optional dict from read_calibration_from_pdf() — overrides OVERLAY_CONFIG geometry.
    """
    cfg = OVERLAY_CONFIG
    pagesize = A4 if cfg["page_size"] == "A4" else (215.9 * mm, 279.4 * mm)
    page_w, page_h = pagesize
    fsize = cfg["font_size"]

    # Use live calibration if provided, else fall back to baked-in config
    def get_page_geom(page_num):
        if calib and page_num in calib:
            p = calib[page_num]
            return (
                [x * mm for x in p["col_x0"]],
                [w * mm for w in p["col_w"]],
                p["data_start_y"],
                p["avg_row_h"],
            )
        elif page_num == 1:
            return (
                [x * mm for x in cfg["col_x_mm"]],
                [w * mm for w in cfg["col_w_mm"]],
                cfg["table_data_start_y_mm"],
                cfg["row_height_mm"],
            )
        else:
            return (
                [x * mm for x in cfg["col_x_mm"]],
                [w * mm for w in cfg["col_w_mm"]],
                cfg["table_data_start_y_mm_page2"],
                cfg["row_height_mm"],
            )

    from collections import defaultdict

    pages = defaultdict(list)
    for entry in rows_to_print:
        pages[entry["page"]].append(entry)
    max_page = max(pages.keys()) if pages else 1

    c = rl_canvas.Canvas(out_path, pagesize=pagesize)
    c.setPageCompression(0)
    c.setFont(OVERLAY_FONT, fsize)

    for page_num in range(1, max_page + 1):
        entries = pages.get(page_num, [])
        col_x, col_w, start_y, rh_mm = get_page_geom(page_num)

        for entry in entries:
            row_idx = entry["target_row"] - 1  # 0-indexed
            vals = entry["values"]  # 7 values; index 5 = Notes
            y_top_mm = start_y + row_idx * rh_mm
            y_center = page_h - (y_top_mm + rh_mm / 2) * mm

            for col_idx, val in enumerate(vals):
                val = str(val)
                if not val or col_idx >= len(col_x):
                    continue

                # Col 5 = Notes: use wrapped, auto-sized, vertically centered drawing
                if col_idx == 5:
                    draw_notes_in_cell(
                        canvas_obj=c,
                        note_text=val,
                        cell_x0_mm=col_x[col_idx] / mm,
                        cell_y_top_mm=y_top_mm,
                        cell_w_mm=col_w[col_idx] / mm,
                        cell_h_mm=rh_mm,
                        page_h_pt=page_h,
                    )
                    continue

                # All other cols: single-line centered
                text_w = c.stringWidth(val, OVERLAY_FONT, fsize)
                x = col_x[col_idx] + (col_w[col_idx] - text_w) / 2
                x = max(
                    col_x[col_idx] + 1 * mm,
                    min(x, col_x[col_idx] + col_w[col_idx] - text_w - 1 * mm),
                )
                c.drawString(x, y_center - (fsize / 2) * 0.352778 * mm, val)

        if page_num < max_page:
            c.showPage()
            c.setFont(OVERLAY_FONT, fsize)

    c.save()
    print(f"  [+] Overlay PDF saved → {out_path}")


# ══════════════════════════════════════════════
# INTERACTIVE ASSET / ROW SELECTOR
# ══════════════════════════════════════════════


def display_asset_table(assets_df, id_col, printed_ids=None):
    """Pretty-print all assets with index numbers."""
    printed_ids = printed_ids or set()
    print()
    print(
        f"  {'#':<4} {'Asset ID':<12} {'Type':<18} {'Brand':<14} {'Model':<16} {'Serial':<18} {'Status'}"
    )
    separator("-", 100)
    for idx, (_, row) in enumerate(assets_df.iterrows(), start=1):
        aid = clean_text(row.get(id_col, ""))
        status = "PRINTED" if str(aid) in printed_ids else "new"
        print(
            f"  {idx:<4} {aid:<12} "
            f"{clean_text(row.get('Asset_Type','')):<18} "
            f"{clean_text(row.get('Brand','')):<14} "
            f"{clean_text(row.get('Model','')):<16} "
            f"{clean_text(row.get('Serial_Number','')):<18} "
            f"{status}"
        )
    separator("-", 100)


def parse_selection(raw, max_n):
    """
    Parse user selection string like "1,3,5-7,9" → set of 0-indexed ints.
    'all' → all indices. Returns sorted list.
    """
    if raw.strip().lower() == "all":
        return list(range(max_n))
    indices = set()
    for part in raw.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            try:
                indices.update(range(int(a) - 1, int(b)))
            except ValueError:
                print(f"  [!] Skipping invalid range: {part}")
        else:
            try:
                indices.add(int(part) - 1)
            except ValueError:
                print(f"  [!] Skipping invalid value: {part}")
    return sorted(i for i in indices if 0 <= i < max_n)


def select_assets_interactively(assets_df, id_col, printed_ids):
    """
    Shows table, lets user pick which assets to include, returns selected df.
    """
    display_asset_table(assets_df, id_col, printed_ids)
    total = len(assets_df)

    print(
        "\n  Select assets to include  (e.g.  1,3,5-7  or  all  or  new  for unprinted only  or  none  to skip)"
    )
    raw = ask("Your selection", default="all")

    if raw.strip().lower() == "none":
        selected = pd.DataFrame()
    elif raw.strip().lower() == "new":
        mask = ~assets_df[id_col].astype(str).isin(printed_ids)
        selected = assets_df[mask].reset_index(drop=True)
    else:
        idx_list = parse_selection(raw, total)
        selected = assets_df.iloc[idx_list].reset_index(drop=True)

    if selected.empty:
        print("  (No table assets selected — you can still add custom items.)")
    else:
        print(f"\n  ✓ {len(selected)} asset(s) selected.")
    return selected


def assign_row_positions(
    selected_df, id_col, notes_cell_w_mm=26.300, notes_cell_h_mm=16.057
):
    """
    For each selected asset, ask PAGE, ROW, and optional NOTES.
    Notes are validated to fit (auto-sized 9→6pt wrapping) before accepting.
    Returns list of {"page": int, "target_row": int (1-indexed), "values": [...], "asset_id": str}
    """
    from reportlab.pdfgen import canvas as _rl_canvas
    from reportlab.lib.pagesizes import A4 as _A4
    import io as _io

    # Temporary canvas just for string width measurements
    _c = _rl_canvas.Canvas(_io.BytesIO(), pagesize=_A4)

    section("ASSIGN PRINT ROWS & NOTES")
    print("  For each asset:")
    print("    1. Enter PAGE and ROW number on the physical paper")
    print("    2. Optionally enter a NOTE (auto-wrapped to fit the Notes cell)")
    print("  Row 1 = first data row on that page table (after header).")
    print("  Page 1 has 7 rows, Page 2 continues with its own rows 1+.")
    print("  Press Enter to accept defaults / skip notes.\n")

    assignments = []
    next_auto_row = 1
    next_auto_page = 1
    padding_mm = 1.5

    for i, (_, row) in enumerate(selected_df.iterrows(), start=1):
        aid = clean_text(row.get(id_col, ""))
        atype = clean_text(row.get("Asset_Type", ""))
        brand = clean_text(row.get("Brand", ""))
        model = clean_text(row.get("Model", ""))
        serial = clean_text(row.get("Serial_Number", ""))

        print(
            f"  ┌─ [{i}/{len(selected_df)}] {aid}  {atype}  {brand} {model}  S/N: {serial}"
        )

        page_num = ask_int(
            "  │  Page number", min_val=1, max_val=2, default=next_auto_page
        )
        row_num = ask_int(
            "  │  Row  number", min_val=1, max_val=50, default=next_auto_row
        )

        # Auto-increment
        next_auto_page = page_num
        if row_num < 7:
            next_auto_row = row_num + 1
        else:
            next_auto_row = 1
            next_auto_page = page_num + 1

        # Notes input with fit validation
        note = ""
        while True:
            raw_note = input(f"  │  Notes (optional, Enter to skip): ").strip()
            if not raw_note:
                note = ""
                break
            # Validate it fits
            fsize, lines = best_fit_note(
                raw_note, OVERLAY_FONT, notes_cell_w_mm, notes_cell_h_mm, padding_mm, _c
            )
            line_h_mm = fsize * 1.2 * 25.4 / 72
            total_h_mm = len(lines) * line_h_mm
            usable_h = notes_cell_h_mm - 2 * padding_mm

            if total_h_mm <= usable_h:
                note = raw_note
                # Show preview of how it will appear
                print(f"  │  ✓ Note fits at {fsize}pt in {len(lines)} line(s):")
                for l in lines:
                    print(f"  │      '{l}'")
                break
            else:
                # This shouldn't normally happen (best_fit_note truncates at 6pt)
                # but show warning and let user shorten
                print(f"  │  ⚠  Text too long even at 6pt. Please shorten it.")
                print(
                    f"  │     Max ~{int((notes_cell_w_mm - 2*padding_mm) / (6*0.6*25.4/72) * (int(usable_h/(6*1.2*25.4/72))))} chars total."
                )

        print(f"  └──────────────────────────────────")
        print()

        values = [
            aid,
            atype,
            brand,
            model,
            serial,
            note,  # Notes cell — wrapped by draw_notes_in_cell at render time
            "",  # Signature — always blank
        ]

        assignments.append(
            {
                "page": page_num,
                "target_row": row_num,
                "values": values,
                "asset_id": aid,
            }
        )

    return assignments


def review_assignments(assignments):
    """Show a final summary before printing."""
    section("REVIEW — OVERLAY ASSIGNMENTS")
    print(
        f"  {'Pg':<4} {'Row':<5} {'Asset ID':<14} {'Type':<16} {'Serial':<18} {'Notes'}"
    )
    separator("-", 90)
    for a in assignments:
        v = a["values"]
        note_preview = (v[5][:30] + "...") if len(v[5]) > 30 else v[5]
        note_display = f'"{note_preview}"' if note_preview else "(none)"
        print(
            f"  {a['page']:<4} {a['target_row']:<5} {v[0]:<14} {v[1]:<16} {v[4]:<18} {note_display}"
        )
    separator("-", 90)
    confirm = ask("Proceed with these assignments?", valid=["y", "n"], default="y")
    return confirm.lower() == "y"


def add_custom_items_interactively(id_col):
    """
    Optionally prompts the user to add custom asset rows not in SharePoint
    (e.g. laptop chargers, bags). Returns a DataFrame with the same column
    structure, or an empty DataFrame if the user skips.
    """
    if ask("Add custom items not in the table?", valid=["y", "n"], default="n") != "y":
        return pd.DataFrame()

    section("CUSTOM ITEMS")
    print("  Enter custom items one by one.")
    print("  Leave Asset ID blank to finish.\n")

    custom_rows = []
    idx = 1
    while True:
        print(f"  ─ Custom item #{idx} ─")
        asset_id = input("    Asset ID    (blank to finish): ").strip()
        if not asset_id:
            break
        asset_type = input("    Item/Type                  : ").strip()
        brand      = input("    Brand                      : ").strip()
        model      = input("    Model                      : ").strip()
        serial     = input("    Serial No.                 : ").strip()
        print()
        custom_rows.append({
            id_col:          asset_id,
            "Asset_Type":    asset_type,
            "Brand":         brand,
            "Model":         model,
            "Serial_Number": serial,
        })
        idx += 1

    if not custom_rows:
        print("  (No custom items added.)\n")
        return pd.DataFrame()

    print(f"  [+] {len(custom_rows)} custom item(s) added.\n")
    return pd.DataFrame(custom_rows)


# ══════════════════════════════════════════════
# FULL DOCUMENT GENERATOR (original flow)
# ══════════════════════════════════════════════


def generate_full_document(
    assets_df,
    id_col,
    emp_name,
    emp_id,
    designation,
    doc_type,
    target_template,
    save_folder,
):
    section("GENERATING FULL DOCUMENT")
    try:
        doc = DocxTemplate(target_template)
    except Exception as e:
        print(f"  [Error] Template missing: {e}")
        return None

    sd = doc.new_subdoc()
    table = sd.add_table(rows=1, cols=7)
    format_table_style(table)

    headers = ["Asset ID", "Item", "Brand", "Model", "Serial No.", "Notes", "Sign"]
    header_row = table.rows[0]
    set_row_height(header_row, 500)
    set_repeat_table_header(header_row)

    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = text
        set_cell_vertical_center(cell)
        for p in cell.paragraphs:
            p.alignment = WD_TABLE_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for _, row in assets_df.iterrows():
        new_row = table.add_row()
        cells = new_row.cells
        set_row_height(new_row, 900)
        keep_row_together(new_row)
        vals = [
            clean_text(row.get(id_col)),
            clean_text(row.get("Asset_Type")),
            clean_text(row.get("Brand")),
            clean_text(row.get("Model")),
            clean_text(row.get("Serial_Number")),
            "",
            "",
        ]
        for i, val in enumerate(vals):
            cells[i].text = val
            set_cell_vertical_center(cells[i])
            for p in cells[i].paragraphs:
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    for _ in range(6):
        empty_row = table.add_row()
        set_row_height(empty_row, 900)
        keep_row_together(empty_row)
        for cell in empty_row.cells:
            set_cell_vertical_center(cell)

    context = {
        "employee_name": emp_name,
        "emp_id": emp_id,
        "designation": designation,
        "document_date": datetime.now().strftime("%d %b %Y"),
        "dynamic_table": sd,
    }

    try:
        doc.render(context)
        clean_name = emp_name.replace(" ", "_")
        out_name = f"{doc_type}_{clean_name}_{emp_id}"
        out_docx = os.path.join(save_folder, f"{out_name}.docx")
        out_pdf = os.path.join(save_folder, f"{out_name}.pdf")

        doc.save(out_docx)
        print(f"  [+] Word saved  → {out_docx}")

        print("  [-] Converting to PDF...")
        convert(out_docx, out_pdf)
        print(f"  [+] PDF saved   → {out_pdf}")
        return out_pdf
    except Exception as e:
        print(f"\n  [Error] Generation failed: {e}")
        return None


# ══════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════


def generate_asset_form():
    import sys

    # ── Special modes ─────────────────────────
    if "--calibrate" in sys.argv:
        generate_calibration_grid()
        return
    if "--recalibrate" in sys.argv:
        idx = sys.argv.index("--recalibrate")
        if idx + 1 < len(sys.argv):
            recalibrate_from_pdf(sys.argv[idx + 1])
        else:
            print("  Usage: python main.py --recalibrate <path_to_form.pdf>")
        return
    if "--config" in sys.argv:
        section("CURRENT OVERLAY_CONFIG")
        for k, v in OVERLAY_CONFIG.items():
            print(f"  {k:<30} = {v}")
        print()
        return

    script_folder = os.path.dirname(os.path.abspath(__file__))
    handover_dir = os.path.join(script_folder, "handover")
    return_dir = os.path.join(script_folder, "return")
    os.makedirs(handover_dir, exist_ok=True)
    os.makedirs(return_dir, exist_ok=True)

    handover_template_path = os.path.join(script_folder, "handover.docx")
    return_template_path = os.path.join(script_folder, "return.docx")

    # ── Welcome ───────────────────────────────
    header("Gravity Asset Form Generator")
    print("  Tips:")
    print(
        "    python main.py --calibrate                    →  print a visual ruler grid PDF"
    )
    print(
        "    python main.py --recalibrate <form.pdf>       →  auto-measure a new form PDF"
    )
    print(
        "    python main.py --config                       →  show current overlay config"
    )
    print()

    # ── Load SharePoint data ──────────────────
    try:
        sheets = get_sharepoint_data(
            SHAREPOINT_FILE_URL,
            SHAREPOINT_TENANT_ID,
            SHAREPOINT_CLIENT_ID,
            SHAREPOINT_CLIENT_SECRET,
            sheet_names=["Employees", "MasterTable"],
        )
        df_emp = sheets["Employees"]
        df_assets = sheets["MasterTable"]
    except Exception as e:
        print(f"  [Error] SharePoint read failed: {e}")
        return

    df_emp.columns = df_emp.columns.str.strip()
    df_assets.columns = df_assets.columns.str.strip()
    df_emp["EmployeeID"] = pd.to_numeric(df_emp["EmployeeID"], errors="coerce")
    df_emp = df_emp.dropna(subset=["EmployeeID"])
    df_emp["EmployeeID"] = df_emp["EmployeeID"].astype(int).astype(str)

    # ── Employee & form type ──────────────────
    try:
        section("EMPLOYEE & FORM SELECTION")
        target_id = ask("Employee ID")
        print()
        print("  Select form type:")
        print("  1. Handover")
        print("  2. Return")
        choice = ask("Choice", valid=["1", "2"])

        if choice == "1":
            target_template = handover_template_path
            doc_type = "Handover"
            save_folder = handover_dir
        else:
            target_template = return_template_path
            doc_type = "Return"
            save_folder = return_dir
    except KeyboardInterrupt:
        print("\n  [Cancelled]")
        return

    # ── Lookup employee ───────────────────────
    emp_row = df_emp[df_emp["EmployeeID"] == target_id]
    if emp_row.empty:
        print(f"\n  [Error] Employee ID '{target_id}' not found.")
        return

    emp_data = emp_row.iloc[0]
    emp_name = clean_text(emp_data.get("FullName", "")).title()
    emp_designation = clean_text(emp_data.get("Designation", "")).title()
    emp_email = clean_text(emp_data.get("Email", "")).lower()
    print(f"\n  [+] Found: {emp_name}  ({emp_designation})")

    # ── Lookup assets ─────────────────────────
    user_col = "Username" if "Username" in df_assets.columns else "Assigned_To"
    id_col = "AssetID" if "AssetID" in df_assets.columns else "Asset_ID"

    df_assets[user_col] = df_assets[user_col].fillna("")
    assets_found = df_assets[df_assets[user_col].str.lower() == emp_email].reset_index(
        drop=True
    )

    if assets_found.empty:
        print(f"  [!] No assets assigned to {emp_email}")
        if ask("Continue with empty table?", valid=["y", "n"], default="n") != "y":
            return

    # ── Print log info ────────────────────────
    printed_ids = get_printed_asset_ids(target_id, doc_type)
    new_count = len(
        [x for x in assets_found[id_col].astype(str) if x not in printed_ids]
    )

    section("PRINT LOG")
    print(f"  Total assets assigned : {len(assets_found)}")
    print(f"  Already printed       : {len(printed_ids)}")
    print(f"  New (unprinted)       : {new_count}")
    print()
    print("  Options:")
    print("  1. Print FULL new document (all assets)")
    print(
        "  2. Print OVERLAY only (new rows on already-printed sheet)  ← recommended for additions"
    )
    print("  3. View print history")
    print("  4. Clear print log for this employee / form type")
    print("  5. Exit")

    mode = ask("Select mode", valid=["1", "2", "3", "4", "5"])

    if mode == "3":
        show_print_log(target_id, doc_type)
        return

    if mode == "4":
        if (
            ask("Are you sure? This will reset the log.", valid=["y", "n"], default="n")
            == "y"
        ):
            clear_print_log(target_id, doc_type)
        return

    if mode == "5":
        return

    # ── MODE 1 — Full document ────────────────
    if mode == "1":
        section("ASSET SELECTION FOR FULL DOCUMENT")
        selected = select_assets_interactively(assets_found, id_col, printed_ids)
        custom_df = add_custom_items_interactively(id_col)
        if not custom_df.empty:
            selected = pd.concat([selected, custom_df], ignore_index=True)
        if selected.empty:
            return

        out_pdf = generate_full_document(
            selected,
            id_col,
            emp_name,
            target_id,
            emp_designation,
            doc_type,
            target_template,
            save_folder,
        )
        if out_pdf:
            asset_ids = selected[id_col].astype(str).tolist()
            if (
                ask(
                    "\n  Mark these assets as printed in log?",
                    valid=["y", "n"],
                    default="y",
                )
                == "y"
            ):
                mark_assets_printed(target_id, doc_type, asset_ids)
            print(f"\n  [SUCCESS] Opening → {out_pdf}")
            os.startfile(out_pdf)

    # ── MODE 2 — Overlay ──────────────────────
    elif mode == "2":
        section("ASSET SELECTION FOR OVERLAY")
        print("  Select which assets to print as an overlay.")
        print("  These will be printed on top of your already-printed form.\n")

        selected = select_assets_interactively(assets_found, id_col, printed_ids)
        custom_df = add_custom_items_interactively(id_col)
        if not custom_df.empty:
            selected = pd.concat([selected, custom_df], ignore_index=True)
        if selected.empty:
            return

        # ── Auto-read calibration from employee's existing PDF ────────────
        clean_name = emp_name.replace(" ", "_")
        existing_pdf = os.path.join(
            save_folder, f"{doc_type}_{clean_name}_{target_id}.pdf"
        )
        calib = None

        if os.path.exists(existing_pdf):
            print(f"\n  [+] Found existing PDF: {os.path.basename(existing_pdf)}")
            print(f"  [-] Reading exact column & row positions from it...")
            calib = read_calibration_from_pdf(existing_pdf)
            if calib:
                p1 = calib.get(1, {})
                print(f"  [✓] Calibration loaded:")
                print(
                    f"      Page 1 → data_start_y={p1.get('data_start_y','?'):.2f}mm  "
                    f"row_h={p1.get('avg_row_h','?'):.3f}mm  "
                    f"cols={len(p1.get('col_x0',[]))}"
                )
                if 2 in calib:
                    p2 = calib[2]
                    print(
                        f"      Page 2 → data_start_y={p2.get('data_start_y','?'):.2f}mm"
                    )
            else:
                print(
                    f"  [!] Could not read calibration from PDF — using baked-in config."
                )
        else:
            print(f"\n  [!] No existing PDF found at: {existing_pdf}")
            print(f"      Using baked-in OVERLAY_CONFIG values.")

        # Row position assignment
        # Pass live Notes cell dimensions if calibration was read from the PDF
        notes_w = (
            calib[1]["col_w"][5]
            if calib and 1 in calib and len(calib[1].get("col_w", [])) > 5
            else 26.300
        )
        notes_h = calib[1]["avg_row_h"] if calib and 1 in calib else 16.057
        assignments = assign_row_positions(
            selected, id_col, notes_cell_w_mm=notes_w, notes_cell_h_mm=notes_h
        )
        if not assignments:
            return

        # Review & confirm
        if not review_assignments(assignments):
            print("  [Cancelled] No file generated.")
            return

        # Generate overlay PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_name = f"Overlay_{doc_type}_{clean_name}_{target_id}_{timestamp}.pdf"
        out_path = os.path.join(save_folder, out_name)

        generate_overlay_pdf(assignments, out_path, calib=calib)

        # Log
        asset_ids = [a["asset_id"] for a in assignments]
        if (
            ask(
                "\n  Mark these assets as printed in log?",
                valid=["y", "n"],
                default="y",
            )
            == "y"
        ):
            mark_assets_printed(target_id, doc_type, asset_ids)

        print(f"\n  [SUCCESS] Overlay ready → {out_path}")
        print(
            "  ⚠  Load the already-printed sheet back into the printer before printing this PDF."
        )
        os.startfile(out_path)


if __name__ == "__main__":
    generate_asset_form()
