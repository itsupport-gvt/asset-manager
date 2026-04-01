"""
Reports route — generates Handover / Return PDFs from the local SQLite DB.
Uses the same DocxTemplate + docx2pdf pipeline as main-v2.py but sources
data from the DB instead of SharePoint, and auto-injects laptop charger rows.
"""
import os
import sys
from pathlib import Path
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from database import get_db
from models_db import DBAsset, DBEmployee

# ── Frozen (PyInstaller) vs dev path resolution ────────────────────────────
_FROZEN  = Path(sys._MEIPASS) if getattr(sys, "frozen", False) else None
_BACKEND = Path(__file__).resolve().parent.parent   # …/backend/
_ROOT    = _BACKEND.parent.parent                   # …/Asset report generator/ (local only)


def _find_template(name: str) -> Path:
    """Look for template in frozen bundle first, then local backend, then repo root."""
    candidates = []
    if _FROZEN:
        candidates.append(_FROZEN / "templates" / name)
    candidates += [_BACKEND / "templates" / name, _BACKEND / name, _ROOT / name]
    for p in candidates:
        if p.exists():
            return p
    return candidates[0]  # will raise FileNotFoundError with a clear path


TEMPLATES = {
    "Handover": lambda: _find_template("handover.docx"),
    "Return":   lambda: _find_template("return.docx"),
}


def _output_dir(doc_type: str) -> Path:
    """Output folder: ASSET_DATA_DIR (Electron) → repo root sibling → backend/output/ (Docker)."""
    asset_data = os.getenv("ASSET_DATA_DIR")
    if asset_data:
        d = Path(asset_data) / "output" / doc_type.lower()
        d.mkdir(parents=True, exist_ok=True)
        return d
    local = _ROOT / doc_type.lower()
    if local.parent.exists() and local.parent != _BACKEND:
        local.mkdir(parents=True, exist_ok=True)
        return local
    d = _BACKEND / "output" / doc_type.lower()
    d.mkdir(parents=True, exist_ok=True)
    return d

router = APIRouter(prefix="/api/report", tags=["reports"])


# ── helpers ────────────────────────────────────────────────────────────

def _asset_rows(assets: list[DBAsset]) -> list[dict]:
    """
    Convert a list of DBAsset objects to flat row dicts, injecting a
    'Laptop Charger' row after every Laptop that has charger data.
    """
    rows = []
    for a in assets:
        rows.append({
            "asset_id":     a.asset_id or "",
            "asset_type":   a.asset_type or "",
            "brand":        a.brand or "",
            "model":        a.model or "",
            "serial_number": a.serial_number or "",
            "notes":        a.notes or "",
            "is_charger":   False,
        })
        # Auto-inject charger row for Laptops
        if a.asset_type == "Laptop" and (a.charger_model or a.charger_serial):
            charger_brand = (a.charger_model or "").split()[0] if (a.charger_model or "").split() else ""
            rows.append({
                "asset_id":     "",
                "asset_type":   "Laptop Charger",
                "brand":        charger_brand,
                "model":        a.charger_model or "",
                "serial_number": a.charger_serial or "",
                "notes":        a.charger_notes or "",
                "is_charger":   True,
            })
    return rows


# ── docx table helpers (adapted from main-v2.py) ───────────────────────

def _build_docx_file(rows: list[dict], emp_name: str, emp_id: str,
                     designation: str, doc_type: str) -> Path:
    """Build and save the .docx file. Returns the saved Path (no PDF conversion)."""
    from docxtpl import DocxTemplate
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    template_path = TEMPLATES[doc_type]()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    output_dir = _output_dir(doc_type)

    clean_name = emp_name.replace(" ", "_")
    out_docx = output_dir / f"{doc_type}_{clean_name}_{emp_id}.docx"
    out_pdf  = output_dir / f"{doc_type}_{clean_name}_{emp_id}.pdf"

    doc = DocxTemplate(str(template_path))
    sd  = doc.new_subdoc()
    table = sd.add_table(rows=1, cols=7)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _set_row_height(row, height=800):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), str(height))
        trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)

    def _keep_together(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        cs.set(qn("w:val"), "true")
        trPr.append(cs)

    def _vcenter(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        v = OxmlElement("w:vAlign")
        v.set(qn("w:val"), "center")
        tcPr.append(v)

    def _repeat_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)

    # Header row
    headers = ["Asset ID", "Item", "Brand", "Model", "Serial No.", "Notes", "Sign"]
    header_row = table.rows[0]
    _set_row_height(header_row, 500)
    _repeat_header(header_row)
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = text
        _vcenter(cell)
        for p in cell.paragraphs:
            p.alignment = WD_TABLE_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_data in rows:
        new_row = table.add_row()
        _set_row_height(new_row, 900)
        _keep_together(new_row)
        vals = [
            row_data["asset_id"],
            row_data["asset_type"],
            row_data["brand"],
            row_data["model"],
            row_data["serial_number"],
            row_data["notes"],
            "",  # Signature blank
        ]
        for i, val in enumerate(vals):
            cell = new_row.cells[i]
            cell.text = val
            _vcenter(cell)
            for p in cell.paragraphs:
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_data.get("is_charger"):
                        run.font.italic = True  # visually distinguish charger rows

    # 6 empty rows for extra signatures
    for _ in range(6):
        empty_row = table.add_row()
        _set_row_height(empty_row, 900)
        _keep_together(empty_row)
        for cell in empty_row.cells:
            _vcenter(cell)

    context = {
        "employee_name":  emp_name,
        "emp_id":         emp_id,
        "designation":    designation,
        "document_date":  datetime.now().strftime("%d %b %Y"),
        "dynamic_table":  sd,
    }
    doc.render(context)
    doc.save(str(out_docx))
    return out_docx


def _build_docx(rows: list[dict], emp_name: str, emp_id: str,
                designation: str, doc_type: str) -> Path:
    """Build .docx then convert to PDF. Returns the PDF Path."""
    from docx2pdf import convert
    out_docx = _build_docx_file(rows, emp_name, emp_id, designation, doc_type)
    out_pdf  = out_docx.with_suffix(".pdf")
    convert(str(out_docx), str(out_pdf))
    return out_pdf


# ── API endpoints ──────────────────────────────────────────────────────

@router.get("/preview/{employee_email}")
def report_preview(employee_email: str, db: Session = Depends(get_db)):
    """Return employee info + pre-computed asset rows (with charger injected)."""
    emp = db.query(DBEmployee).filter(DBEmployee.email == employee_email).first()
    if not emp:
        raise HTTPException(status_code=404, detail=f"Employee '{employee_email}' not found")

    assets = (
        db.query(DBAsset)
        .filter(
            DBAsset.assigned_to_email == employee_email,
            DBAsset.status != "Retired",
        )
        .order_by(DBAsset.asset_type, DBAsset.asset_id)
        .all()
    )

    return {
        "employee": {
            "email":       emp.email,
            "full_name":   emp.full_name,
            "employee_id": emp.employee_id or "",
            "designation": emp.designation or "",
            "display":     emp.employee_display or emp.full_name,
        },
        "rows": _asset_rows(assets),
        "asset_count": len(assets),
    }


@router.post("/generate")
def generate_report(body: dict, db: Session = Depends(get_db)):
    """
    Generate a Handover or Return PDF for an employee.

    Body: { "employee_email": str, "doc_type": "Handover"|"Return", "excluded_ids": [str] }
    Returns the PDF as a file download.
    """
    employee_email = body.get("employee_email", "")
    doc_type       = body.get("doc_type", "Handover")
    excluded_ids   = set(body.get("excluded_ids", []))

    if doc_type not in TEMPLATES:
        raise HTTPException(status_code=400, detail="doc_type must be 'Handover' or 'Return'")

    emp = db.query(DBEmployee).filter(DBEmployee.email == employee_email).first()
    if not emp:
        raise HTTPException(status_code=404, detail=f"Employee '{employee_email}' not found")

    assets = (
        db.query(DBAsset)
        .filter(
            DBAsset.assigned_to_email == employee_email,
            DBAsset.status != "Retired",
        )
        .order_by(DBAsset.asset_type, DBAsset.asset_id)
        .all()
    )

    # Build rows with charger injection, then filter excluded
    all_rows = _asset_rows(assets)
    rows = [r for r in all_rows if r["asset_id"] not in excluded_ids]

    if not rows:
        raise HTTPException(status_code=400, detail="No assets selected for the report")

    emp_name    = (emp.full_name or "Unknown").title()
    emp_id      = emp.employee_id or ""
    designation = (emp.designation or "").title()

    try:
        pdf_path = _build_docx(rows, emp_name, emp_id, designation, doc_type)
    except FileNotFoundError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")

    filename = f"{doc_type}_{emp_name.replace(' ', '_')}_{emp_id}.pdf"
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/generate-docx")
def generate_report_docx(body: dict, db: Session = Depends(get_db)):
    """
    Generate a Handover or Return Word document (.docx) for an employee.

    Body: { "employee_email": str, "doc_type": "Handover"|"Return", "excluded_ids": [str] }
    Returns the .docx as a file download (no PDF conversion required).
    """
    employee_email = body.get("employee_email", "")
    doc_type       = body.get("doc_type", "Handover")
    excluded_ids   = set(body.get("excluded_ids", []))

    if doc_type not in TEMPLATES:
        raise HTTPException(status_code=400, detail="doc_type must be 'Handover' or 'Return'")

    emp = db.query(DBEmployee).filter(DBEmployee.email == employee_email).first()
    if not emp:
        raise HTTPException(status_code=404, detail=f"Employee '{employee_email}' not found")

    assets = (
        db.query(DBAsset)
        .filter(
            DBAsset.assigned_to_email == employee_email,
            DBAsset.status != "Retired",
        )
        .order_by(DBAsset.asset_type, DBAsset.asset_id)
        .all()
    )

    all_rows = _asset_rows(assets)
    rows = [r for r in all_rows if r["asset_id"] not in excluded_ids]

    if not rows:
        raise HTTPException(status_code=400, detail="No assets selected for the report")

    emp_name    = (emp.full_name or "Unknown").title()
    emp_id      = emp.employee_id or ""
    designation = (emp.designation or "").title()

    try:
        docx_path = _build_docx_file(rows, emp_name, emp_id, designation, doc_type)
    except FileNotFoundError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word generation failed: {e}")

    filename = f"{doc_type}_{emp_name.replace(' ', '_')}_{emp_id}.docx"
    return FileResponse(
        str(docx_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
